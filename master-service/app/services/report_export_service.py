from __future__ import annotations

import csv
from dataclasses import dataclass
from datetime import datetime, timezone
from io import BytesIO, StringIO

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from sqlalchemy import and_
from sqlalchemy.orm import Session, selectinload

from app.database.models import (
    PlacementRecord,
    Student,
    TpoAnalysisGroup,
    TpoAnalysisGroupMember,
    TpoGroupRound,
    TpoGroupRoundMember,
    TpoSettings,
)


@dataclass
class TpoOverviewSummary:
    total_students: int
    unplaced_eligible_students: int
    active_groups: int
    placed_students: int
    internships_count: int


@dataclass
class TpoGroupSummaryRow:
    group_id: int
    title: str
    company_name: str
    role_type: str
    total_members: int
    current_round_no: int
    total_rounds: int
    qualified_count: int
    rejected_count: int
    pending_count: int
    round_state: str
    created_at: datetime


@dataclass
class TpoPlacementRow:
    student_id: int
    name: str
    email: str
    roll_no: str
    branch: str
    company_name: str
    offer_type: str
    pay_amount: float | None
    updated_at: datetime


@dataclass
class TpoReportPayload:
    generated_at: datetime
    generated_by: str
    institute_name: str
    overview: TpoOverviewSummary
    groups: list[TpoGroupSummaryRow]
    placements: list[TpoPlacementRow]


class TpoReportExportService:
    def __init__(self, db: Session) -> None:
        self.db = db

    def build_payload(
        self,
        *,
        tpo_username: str,
        group_id: int | None = None,
        branch: str | None = None,
        placed_only: bool = False,
        date_from: datetime | None = None,
        date_to: datetime | None = None,
    ) -> TpoReportPayload:
        settings = self.db.query(TpoSettings).filter(TpoSettings.tpo_username == tpo_username).one_or_none()
        institute_name = (settings.institute_name if settings else None) or "Institute"

        total_students = self.db.query(Student).count()
        active_groups = self.db.query(TpoAnalysisGroup).count()
        placed_students = (
            self.db.query(PlacementRecord.student_id)
            .filter(PlacementRecord.is_active.is_(True))
            .distinct()
            .count()
        )
        internships_count = (
            self.db.query(PlacementRecord.student_id)
            .filter(
                PlacementRecord.is_active.is_(True),
                PlacementRecord.offer_type.ilike("internship"),
            )
            .distinct()
            .count()
        )
        unplaced_eligible_students = (
            self.db.query(Student.id)
            .outerjoin(
                PlacementRecord,
                and_(PlacementRecord.student_id == Student.id, PlacementRecord.is_active.is_(True)),
            )
            .filter(PlacementRecord.id.is_(None), Student.has_active_backlog.is_(False))
            .count()
        )
        overview = TpoOverviewSummary(
            total_students=total_students,
            unplaced_eligible_students=unplaced_eligible_students,
            active_groups=active_groups,
            placed_students=placed_students,
            internships_count=internships_count,
        )

        groups_query = self.db.query(TpoAnalysisGroup).options(
            selectinload(TpoAnalysisGroup.members).selectinload(TpoAnalysisGroupMember.student)
        )
        if group_id is not None:
            groups_query = groups_query.filter(TpoAnalysisGroup.id == group_id)
        groups = groups_query.order_by(TpoAnalysisGroup.created_at.desc(), TpoAnalysisGroup.id.desc()).all()
        group_rows = [self._to_group_row(group) for group in groups]

        placements_query = (
            self.db.query(PlacementRecord, Student)
            .join(Student, Student.id == PlacementRecord.student_id)
            .filter(PlacementRecord.is_active.is_(True))
        )
        if placed_only:
            placements_query = placements_query.filter(PlacementRecord.is_active.is_(True))
        if date_from is not None:
            placements_query = placements_query.filter(PlacementRecord.updated_at >= date_from)
        if date_to is not None:
            placements_query = placements_query.filter(PlacementRecord.updated_at <= date_to)
        if branch is not None and branch.strip():
            placements_query = placements_query.filter(Student.branch.ilike(branch.strip()))
        if group_id is not None:
            member_ids = {
                member.student_id
                for group in groups
                for member in group.members
                if member.student_id is not None
            }
            if member_ids:
                placements_query = placements_query.filter(PlacementRecord.student_id.in_(member_ids))
            else:
                placements_query = placements_query.filter(PlacementRecord.student_id == -1)
        placement_rows = [
            TpoPlacementRow(
                student_id=student.id,
                name=student.name,
                email=student.email,
                roll_no=student.roll_no or "-",
                branch=student.branch,
                company_name=record.company_name,
                offer_type=record.offer_type,
                pay_amount=record.pay_amount,
                updated_at=record.updated_at,
            )
            for record, student in placements_query.order_by(PlacementRecord.updated_at.desc()).all()
        ]

        return TpoReportPayload(
            generated_at=datetime.now(timezone.utc),
            generated_by=tpo_username,
            institute_name=institute_name,
            overview=overview,
            groups=group_rows,
            placements=placement_rows,
        )

    def export_csv(self, payload: TpoReportPayload) -> bytes:
        out = StringIO()
        writer = csv.writer(out)
        writer.writerow(
            [
                "student_id",
                "name",
                "email",
                "roll_no",
                "branch",
                "company_name",
                "offer_type",
                "pay_amount",
                "updated_at",
            ]
        )
        for row in payload.placements:
            writer.writerow(
                [
                    row.student_id,
                    row.name,
                    row.email,
                    row.roll_no,
                    row.branch,
                    row.company_name,
                    row.offer_type,
                    row.pay_amount if row.pay_amount is not None else "",
                    row.updated_at.isoformat(),
                ]
            )
        return out.getvalue().encode("utf-8")

    def export_xlsx(self, payload: TpoReportPayload) -> bytes:
        wb = Workbook()
        overview_ws = wb.active
        overview_ws.title = "Overview"
        overview_ws.append(["Metric", "Value"])
        overview_ws.append(["Total Students", payload.overview.total_students])
        overview_ws.append(["Unplaced Eligible Students", payload.overview.unplaced_eligible_students])
        overview_ws.append(["Active Groups", payload.overview.active_groups])
        overview_ws.append(["Placed Students", payload.overview.placed_students])
        overview_ws.append(["Internships", payload.overview.internships_count])
        overview_ws.append(["Generated By", payload.generated_by])
        overview_ws.append(["Generated At", payload.generated_at.isoformat()])

        groups_ws = wb.create_sheet("Groups")
        groups_ws.append(
            [
                "group_id",
                "title",
                "company_name",
                "role_type",
                "total_members",
                "current_round_no",
                "total_rounds",
                "qualified_count",
                "rejected_count",
                "pending_count",
                "round_state",
                "created_at",
            ]
        )
        for row in payload.groups:
            groups_ws.append(
                [
                    row.group_id,
                    row.title,
                    row.company_name,
                    row.role_type,
                    row.total_members,
                    row.current_round_no,
                    row.total_rounds,
                    row.qualified_count,
                    row.rejected_count,
                    row.pending_count,
                    row.round_state,
                    row.created_at.isoformat(),
                ]
            )

        placements_ws = wb.create_sheet("Placements")
        placements_ws.append(
            [
                "student_id",
                "name",
                "email",
                "roll_no",
                "branch",
                "company_name",
                "offer_type",
                "pay_amount",
                "updated_at",
            ]
        )
        for row in payload.placements:
            placements_ws.append(
                [
                    row.student_id,
                    row.name,
                    row.email,
                    row.roll_no,
                    row.branch,
                    row.company_name,
                    row.offer_type,
                    row.pay_amount,
                    row.updated_at.isoformat(),
                ]
            )

        stream = BytesIO()
        wb.save(stream)
        return stream.getvalue()

    def export_docx(self, payload: TpoReportPayload) -> bytes:
        doc = Document()
        doc.add_heading("TPO Placement Report", level=1)
        doc.add_paragraph(f"Institution: {payload.institute_name}")
        doc.add_paragraph(f"Generated By: {payload.generated_by}")
        doc.add_paragraph(f"Generated At: {payload.generated_at.isoformat()}")

        doc.add_heading("Overview", level=2)
        doc.add_paragraph(f"Total Students: {payload.overview.total_students}")
        doc.add_paragraph(f"Unplaced Eligible Students: {payload.overview.unplaced_eligible_students}")
        doc.add_paragraph(f"Active Groups: {payload.overview.active_groups}")
        doc.add_paragraph(f"Placed Students: {payload.overview.placed_students}")
        doc.add_paragraph(f"Internships: {payload.overview.internships_count}")

        doc.add_heading("Group Summary", level=2)
        groups_table = doc.add_table(rows=1, cols=6)
        groups_table.style = "Table Grid"
        headers = groups_table.rows[0].cells
        headers[0].text = "Group"
        headers[1].text = "Company"
        headers[2].text = "Members"
        headers[3].text = "Round"
        headers[4].text = "Qualified/Rejected/Pending"
        headers[5].text = "State"
        for row in payload.groups:
            cells = groups_table.add_row().cells
            cells[0].text = row.title
            cells[1].text = row.company_name
            cells[2].text = str(row.total_members)
            cells[3].text = f"{row.current_round_no}/{row.total_rounds}"
            cells[4].text = f"{row.qualified_count}/{row.rejected_count}/{row.pending_count}"
            cells[5].text = row.round_state

        doc.add_heading("Placements", level=2)
        placement_table = doc.add_table(rows=1, cols=7)
        placement_table.style = "Table Grid"
        ph = placement_table.rows[0].cells
        ph[0].text = "Name"
        ph[1].text = "Roll No"
        ph[2].text = "Branch"
        ph[3].text = "Company"
        ph[4].text = "Offer Type"
        ph[5].text = "Pay"
        ph[6].text = "Updated At"
        for row in payload.placements:
            cells = placement_table.add_row().cells
            cells[0].text = row.name
            cells[1].text = row.roll_no
            cells[2].text = row.branch
            cells[3].text = row.company_name
            cells[4].text = row.offer_type
            cells[5].text = str(row.pay_amount if row.pay_amount is not None else "")
            cells[6].text = row.updated_at.isoformat()

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

        stream = BytesIO()
        doc.save(stream)
        return stream.getvalue()

    def export_pdf(self, payload: TpoReportPayload) -> bytes:
        stream = BytesIO()
        doc = SimpleDocTemplate(stream, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
        styles = getSampleStyleSheet()
        story = []
        story.append(Paragraph("TPO Placement Report", styles["Title"]))
        story.append(Spacer(1, 8))
        story.append(Paragraph(f"Institution: {payload.institute_name}", styles["Normal"]))
        story.append(Paragraph(f"Generated By: {payload.generated_by}", styles["Normal"]))
        story.append(Paragraph(f"Generated At: {payload.generated_at.isoformat()}", styles["Normal"]))
        story.append(Spacer(1, 12))

        story.append(Paragraph("Overview", styles["Heading2"]))
        story.append(
            Paragraph(
                (
                    f"Total Students: {payload.overview.total_students} | "
                    f"Unplaced Eligible: {payload.overview.unplaced_eligible_students} | "
                    f"Active Groups: {payload.overview.active_groups} | "
                    f"Placed Students: {payload.overview.placed_students} | "
                    f"Internships: {payload.overview.internships_count}"
                ),
                styles["Normal"],
            )
        )
        story.append(Spacer(1, 10))

        story.append(Paragraph("Group Summary", styles["Heading2"]))
        group_data = [["Group", "Company", "Members", "Round", "Q/R/P", "State"]]
        for row in payload.groups:
            group_data.append(
                [
                    row.title[:30],
                    row.company_name[:22],
                    str(row.total_members),
                    f"{row.current_round_no}/{row.total_rounds}",
                    f"{row.qualified_count}/{row.rejected_count}/{row.pending_count}",
                    row.round_state,
                ]
            )
        group_table = Table(group_data, repeatRows=1)
        group_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E2E8F0")),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ]
            )
        )
        story.append(group_table)
        story.append(Spacer(1, 10))

        story.append(Paragraph("Placement Records", styles["Heading2"]))
        placement_data = [["Name", "Roll", "Branch", "Company", "Offer", "Pay"]]
        for row in payload.placements[:150]:
            placement_data.append(
                [
                    row.name[:22],
                    row.roll_no,
                    row.branch,
                    row.company_name[:22],
                    row.offer_type,
                    str(row.pay_amount if row.pay_amount is not None else ""),
                ]
            )
        placement_table = Table(placement_data, repeatRows=1)
        placement_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E2E8F0")),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ]
            )
        )
        story.append(placement_table)
        if len(payload.placements) > 150:
            story.append(Spacer(1, 6))
            story.append(Paragraph("Note: Only the first 150 placement rows are shown in PDF table view.", styles["Italic"]))

        doc.build(story)
        return stream.getvalue()

    def _to_group_row(self, group: TpoAnalysisGroup) -> TpoGroupSummaryRow:
        current_round = (
            self.db.query(TpoGroupRound)
            .filter(TpoGroupRound.group_id == group.id, TpoGroupRound.round_no == group.current_round_no)
            .one_or_none()
        )
        qualified_count = 0
        rejected_count = 0
        pending_count = 0
        if current_round is not None:
            members = self.db.query(TpoGroupRoundMember).filter(TpoGroupRoundMember.round_id == current_round.id).all()
            for member in members:
                if member.status == "qualified":
                    qualified_count += 1
                elif member.status == "rejected":
                    rejected_count += 1
                else:
                    pending_count += 1

        return TpoGroupSummaryRow(
            group_id=group.id,
            title=group.title,
            company_name=group.company_name or "-",
            role_type=group.role_type or "-",
            total_members=len(group.members),
            current_round_no=group.current_round_no,
            total_rounds=group.total_rounds,
            qualified_count=qualified_count,
            rejected_count=rejected_count,
            pending_count=pending_count,
            round_state=group.round_state,
            created_at=group.created_at,
        )
